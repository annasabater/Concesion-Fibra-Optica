"""oferta_tecnica_docx.py — memoria descriptiva en DOCX para TAREA 5.

Genera un documento Word más extenso que el PDF (~15-25 páginas) con
memoria, marco normativo, metodología, resultados y conclusiones.
Pensado para acompañar la presentación oral del 12/05/2026.

Uso:
    python -m src.oferta_tecnica_docx                      # base (default)
    python -m src.oferta_tecnica_docx --escenario optimista
"""

from __future__ import annotations

import argparse
import logging
from datetime import date
from pathlib import Path

import pandas as pd
import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT_DIR = ROOT / "outputs"
DEFAULT_DATA_DIR = ROOT / "data"
ESCENARIOS = ("pesimista", "base", "optimista")

PRIMARY = RGBColor(0x1A, 0x54, 0x90)
SECONDARY = RGBColor(0x21, 0x61, 0x8C)
GREY = RGBColor(0x5D, 0x6D, 0x7E)
ACCENT_BG = "FEF9E7"


# ---------------------------------------------------------------------------
# Carga de datos
# ---------------------------------------------------------------------------

def _load_outputs(out_root: Path, esc: str) -> dict:
    base = out_root / f"escenario_{esc}"
    return {
        "traffic": pd.read_csv(base / "traffic_per_muni.csv"),
        "viability": pd.read_csv(base / "municipal_viability.csv"),
        "deployment": pd.read_csv(base / "deployment_priority.csv"),
        "equipment": pd.read_csv(base / "equipment_per_node.csv"),
        "stats_ring": pd.read_csv(base / "stats_by_ring.csv"),
    }


def _load_decisiones(path: Path, esc: str) -> dict:
    with path.open("r", encoding="utf-8") as f:
        raw = yaml.safe_load(f) or {}
    base = raw.get("escenarios", {}).get("base", {})
    override = raw.get("escenarios", {}).get(esc, {})
    return _deep_merge(base, override)


def _deep_merge(base: dict, override: dict) -> dict:
    out = dict(base)
    for k, v in override.items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


# ---------------------------------------------------------------------------
# Helpers de formato
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level <= 1 else SECONDARY


def _add_paragraph(
    doc: Document, text: str, bold: bool = False, italic: bool = False,
    size: int = 11, justify: bool = True,
) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _add_callout(doc: Document, text: str) -> None:
    """Recuadro destacado para hallazgos clave."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, ACCENT_BG)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(10)
    # Borde fino dorado
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), "F1C40F")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_table(
    doc: Document, header: list[str], rows: list[list],
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabecera
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1A5490")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cuerpo
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = para.add_run(str(value))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)


# ---------------------------------------------------------------------------
# Bloques del documento
# ---------------------------------------------------------------------------

def _block_cover(doc: Document, escenario: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proyecto ABAST")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Concesión a 20 años de red de fibra óptica")
    r.font.size = Pt(14)
    r.font.color.rgb = SECONDARY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Memoria de la oferta técnica — TAREA 5")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Equipos y solución técnica")
    r.font.size = Pt(13)
    r.font.color.rgb = GREY

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Escenario defendido: {escenario}")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fecha de entrega: 12 de mayo de 2026")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Fecha del documento: {date.today().isoformat()}")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Master en Ingeniería de Telecomunicaciones — La Salle URL\n"
        "Projectes de Xarxes i Sistemes de Telecomunicació II"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    doc.add_page_break()


def _block_intro(doc: Document) -> None:
    _add_heading(doc, "Resumen ejecutivo", level=1)
    _add_paragraph(
        doc,
        "El presente documento constituye la memoria descriptiva de la "
        "oferta técnica para la concesión a 20 años de la red de fibra "
        "óptica del territorio ABAST. Se enmarca dentro de la TAREA 5 de "
        "la asignatura Projectes de Xarxes i Sistemes de Telecomunicació II "
        "del Master en Ingeniería de Telecomunicaciones de La Salle URL, y "
        "se presentará el 12 de mayo de 2026.",
    )
    _add_paragraph(
        doc,
        "El procedimiento se desarrolla bajo el modelo de diálogo "
        "competitivo de la Ley de Contratos del Sector Público (LCSP), "
        "con tres entregas escalonadas (oferta técnica el 26/05, oferta "
        "económica el 09/06, oferta final el 30/06) y una ronda de "
        "negociación intermedia. La adjudicataria construirá y explotará "
        "una red de acceso de 4.825 km que conecta los 900 municipios y "
        "las 5.000 sedes públicas del territorio, sobre una capa troncal "
        "y de agregación ya construidas (1.000 km y 2.000 km "
        "respectivamente, con un nodo central A900 sobre los 3 anillos "
        "troncales y los 2 primeros anillos de agregación, que actúa como "
        "datacenter del operador).",
    )
    _add_paragraph(
        doc,
        "El modelo de negocio combina dos fuentes de ingreso: "
        "autoprestación a la administración (1.000 €/mes por sede pública, "
        "alta gratuita) y servicio mayorista a otros operadores presentes "
        "(700 €/mes recurrentes + 1.500 € de alta por sede). La cobertura "
        "del 100% es obligatoria, lo que limita las decisiones a la "
        "estrategia de captura de cuota mayorista, el dimensionamiento "
        "técnico y el plan plurianual de despliegue.",
    )
    _add_paragraph(
        doc,
        "Esta memoria se estructura en cinco bloques: dimensionamiento de "
        "tráfico (§1), selección de equipos por nodo (§2), justificación "
        "de las decisiones técnicas defendidas (§3), análisis "
        "geomarketing de viabilidad y orden de despliegue por municipio "
        "(§4) y hallazgos clave de la solución (§5). Los resultados se "
        "presentan en escenario base, con tablas comparativas frente a "
        "los escenarios optimista y pesimista que sostienen el equilibrio "
        "económico ante el mecanismo de reequilibrio de la LCSP.",
    )


def _block_1_traffic(
    doc: Document, scenarios_data: dict, chart_path: Path,
) -> None:
    _add_heading(doc, "1. Dimensionamiento del tráfico", level=1)

    _add_heading(doc, "1.1. Marco metodológico", level=2)
    _add_paragraph(
        doc,
        "El profesor formalizó en clase del 07/05/2026 la fórmula de "
        "tráfico mayorista, separándola en dos factores independientes "
        "que antes estaban mezclados en un único parámetro α: la cuota "
        "entre operadores (cómo nos repartimos el mercado con los "
        "operadores ya presentes en cada municipio) y la penetración de "
        "fibra (qué porcentaje de hogares contrata realmente fibra, "
        "habida cuenta de que parte del parque sigue con cable, ADSL o "
        "sólo móvil). La separación es clave porque permite distinguir "
        "factor exógeno (penetración, dependiente del mercado) de "
        "factor endógeno (cuota, dependiente de nuestra estrategia "
        "comercial), lo que es decisivo para la negociación del "
        "reequilibrio económico.",
    )
    _add_paragraph(doc, "Las fórmulas aplicadas son:", bold=True)
    _add_paragraph(
        doc,
        "Tráfico ABAST por municipio = sedes_abast × bw_por_sede / "
        "overbooking_abast. En nuestro escenario base, bw_por_sede = "
        "100 Mbps y overbooking = 1:1 (sin overbooking, BW garantizado a "
        "los servicios públicos críticos).",
    )
    _add_paragraph(
        doc,
        "Tráfico mayorista por municipio en el año t = "
        "(hab / habitantes_por_hogar) × cuota_operadores(t) × "
        "penetracion_fibra × bw_por_hogar / overbooking. En el escenario "
        "base: hab/2,5 (ratio INE), cuota con estrategia rampa que "
        "alcanza α=0,40 al año 8, penetración del 70% (mix urbano/rural), "
        "bw=100 Mbps y overbooking 1:20 residencial estándar de la "
        "industria.",
    )

    _add_heading(doc, "1.2. Resultados territoriales en año 20", level=2)
    rows = []
    for esc in ESCENARIOS:
        t = scenarios_data[esc]["traffic"]
        e = scenarios_data[esc]["equipment"]
        bw_a900 = float(e.loc[e["municipio"] == "A900", "bw_acumulado_mbps"].iloc[0])
        rows.append([
            esc.capitalize(),
            f"{t['bw_abast_mbps'].sum() / 1000:,.1f}",
            f"{t['bw_mayorista_mbps'].sum() / 1000:,.1f}",
            f"{t['bw_total_mbps'].sum() / 1000:,.1f}",
            f"{bw_a900 / 1000:,.1f}",
        ])
    _add_table(
        doc,
        ["Escenario", "ABAST (Gbps)", "Mayorista (Gbps)",
         "Total (Gbps)", "BW en A900 (Gbps)"],
        rows,
        [3.5, 2.8, 3.2, 2.8, 3.2],
    )
    _add_paragraph(doc, "")

    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figura 1.1 — Tráfico territorial año 20 por escenario")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    _add_heading(doc, "1.3. Validación y sanity checks", level=2)
    _add_paragraph(
        doc,
        "El tráfico ABAST permanece constante en los tres escenarios "
        "(500 Gbps = 5.000 sedes × 100 Mbps), lo que es coherente con "
        "que la autoprestación es determinística (sin decisiones de "
        "mercado) y la cobertura está fijada al 100%. La diferencia "
        "entre escenarios viene dada exclusivamente por el tráfico "
        "mayorista, que escala con el producto α × penetración: 0,165 "
        "(pesimista), 0,280 (base) y 0,425 (optimista). El BW acumulado "
        "en A900 iguala la suma del tráfico territorial en cada "
        "escenario, lo que valida la correctitud del recorrido en árbol "
        "del módulo equipment.py.",
    )


def _block_2_equipment(
    doc: Document, df_equipment: pd.DataFrame, chart_path: Path,
) -> None:
    _add_heading(doc, "2. Selección de equipos por nodo", level=1)

    _add_heading(doc, "2.1. Algoritmo de cálculo", level=2)
    _add_paragraph(
        doc,
        "La selección de equipos se realiza en dos pasadas. La primera "
        "calcula el tráfico acumulado por nodo recorriendo el grafo de "
        "la topología en orden topológico desde las hojas del árbol de "
        "acceso hacia A900: cada nodo absorbe el tráfico de sus "
        "predecesores en el árbol. A continuación, cada anillo de "
        "agregación reenvía su total al gateway (nodo del anillo que "
        "también pertenece a un anillo troncal), y cada anillo troncal "
        "reenvía a A900. La segunda pasada elige, para cada nodo, el "
        "equipo más pequeño que satisface simultáneamente dos criterios: "
        "el número de puertos requeridos (n_links contra los umbrales "
        "8/16/36 de decisiones.yaml) y la capacidad nominal en Mbps "
        "(10/20/40 Gbps para los Ethernet de agregación, 100 Gbps para "
        "MPLS, 400 Gbps para óptico 40λ).",
    )
    _add_paragraph(
        doc,
        "Cuando el tráfico excede la capacidad de la opción seleccionada "
        "por puertos, se escala en cascada al siguiente equipo. Si supera "
        "los 40 Gbps de un único 40p, se despliega MPLS troncal en "
        "paralelo con N cajas 40p. Si supera los 100 Gbps de un MPLS "
        "típico, se añade óptico 40λ.",
    )

    _add_heading(doc, "2.2. Distribución de equipos en el escenario base", level=2)
    dist = (
        df_equipment.groupby(["tier", "equipo_principal"])
        .size().reset_index(name="n_nodos")
        .sort_values(["tier", "n_nodos"], ascending=[True, False])
    )
    rows = [
        [r["tier"].capitalize(), r["equipo_principal"], int(r["n_nodos"])]
        for _, r in dist.iterrows()
    ]
    _add_table(doc, ["Tier", "Equipo principal", "Nº de nodos"], rows,
               [3.0, 9.0, 3.0])
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "La práctica totalidad de los 799 nodos de acceso queda cubierta "
        "con equipo cliente estándar (300 €/sede). En la capa de "
        "agregación, la elección es heterogénea: 10 nodos con 10p, 51 "
        "con 20p y 27 con 40p, lo que refleja la concentración del "
        "tráfico en unos pocos nodos críticos. Tres nodos de agregación "
        "(A800 y dos más) requieren MPLS combinado con varias cajas 40p "
        "y óptico 40λ por encontrarse en la cadena de la espina dorsal "
        "A2. Los 10 nodos troncales reciben todos MPLS + óptico 40λ "
        "como equipamiento estándar, en línea con su rol de backbone.",
    )

    _add_heading(doc, "2.3. CAPEX de equipos", level=2)
    capex_cliente = float(df_equipment["equipo_cliente_capex"].sum())
    capex_agreg = float(
        df_equipment.loc[df_equipment["tier"] == "agregacion", "capex_equipo"].sum()
    )
    capex_troncal = float(
        df_equipment.loc[df_equipment["tier"] == "troncal", "capex_equipo"].sum()
    )
    capex_chasis = float(df_equipment["capex_chasis"].sum())
    capex_dc = float(df_equipment["capex_extra"].sum())
    capex_total = capex_cliente + capex_agreg + capex_troncal + capex_chasis + capex_dc

    capex_rows = [
        ["Equipo cliente (300 €/sede)", f"{capex_cliente:,.0f} €",
         f"{capex_cliente / capex_total * 100:.1f}%"],
        ["Equipo agregación (10p/20p/40p/MPLS)", f"{capex_agreg:,.0f} €",
         f"{capex_agreg / capex_total * 100:.1f}%"],
        ["Equipo troncal (MPLS + óptico 40λ)", f"{capex_troncal:,.0f} €",
         f"{capex_troncal / capex_total * 100:.1f}%"],
        ["Chasis nodos (100 k€/nodo)", f"{capex_chasis:,.0f} €",
         f"{capex_chasis / capex_total * 100:.1f}%"],
        ["Datacenter A900 (extra)", f"{capex_dc:,.0f} €",
         f"{capex_dc / capex_total * 100:.1f}%"],
        ["TOTAL", f"{capex_total:,.0f} €", "100,0%"],
    ]
    _add_table(doc, ["Categoría", "CAPEX", "% del total"], capex_rows,
               [8.0, 4.5, 2.5])
    _add_paragraph(doc, "")
    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figura 2.1 — CAPEX de equipos por escenario")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
    _add_paragraph(
        doc,
        "El CAPEX de equipos en el escenario base asciende a 18,98 M€, "
        "y varía menos de 1 M€ entre los tres escenarios (18,76 M€ "
        "pesimista, 19,39 M€ optimista). La razón es que la capa troncal "
        "y de agregación está dimensionada por capacidad y cobertura, no "
        "por demanda: los nodos requieren chasis y MPLS aunque el "
        "tráfico mayorista varíe. La sensibilidad al escenario se "
        "manifestará principalmente en la cuenta de resultados (TAREA 7), "
        "no en el CAPEX inicial.",
    )


def _block_3_decisions(doc: Document, decisiones: dict) -> None:
    _add_heading(doc, "3. Decisiones técnicas defendidas", level=1)
    _add_paragraph(
        doc,
        "Todas las palancas variables de la oferta viven en el archivo "
        "data/decisiones.yaml, lo que permite recalcular la pipeline "
        "completa en menos de 30 segundos ante cualquier cambio. Esta "
        "sección recoge las 14 decisiones más relevantes del escenario "
        "base, con su justificación técnica o comercial.",
    )

    may = decisiones["mayorista"]
    abast = decisiones["abast"]
    eq = decisiones["equipos"]
    desp = decisiones["despliegue"]
    infra = decisiones["infraestructura_acceso"]
    dc = decisiones["datacenter_a900"]

    rows = [
        ["Estrategia cuota mayorista", may["cuota_mercado"]["tipo"],
         "Rampa α(t): captura realista, no plena el año 1"],
        ["α objetivo", f"{may['cuota_mercado']['alpha_objetivo']:.2f}",
         "Compromiso entre conservador y agresivo"],
        ["Años de rampa", str(may["cuota_mercado"]["anios_rampa"]),
         "Captura comercial progresiva en 8 años"],
        ["Penetración fibra residencial",
         f"{may['penetracion_fibra_pct'] * 100:.0f}%",
         "Mix urbano/rural; pizarra del 07/05/2026"],
        ["BW comercial por hogar",
         f"{may['bw_por_hogar_mbps']} Mbps",
         "Producto wholesale residencial estándar"],
        ["Overbooking residencial", f"1:{int(may['overbooking'])}",
         "Estándar industria; no degrada SLA"],
        ["Overbooking ABAST", f"1:{int(abast['overbooking'])}",
         "Sin overbooking; BW garantizado a sedes públicas"],
        ["Umbrales 10p/20p/40p",
         f"{eq['umbral_10p_a_20p']}/{eq['umbral_20p_a_40p']}/{eq['umbral_40p_a_mpls']}",
         "Escalado en cascada por puertos y capacidad"],
        ["Redundancia equipos", eq["redundancia"],
         "1+1 sólo en troncal y A900: balance CAPEX/SLA"],
        ["Estrategia despliegue", desp["estrategia"],
         "Espina dorsal y ciudades primero, expansión radial"],
        ["Tope km/año", f"{desp['km_max_anuales']} km",
         "Cuello de botella obra civil + RRHH"],
        ["Punto de acceso", infra["tipo_punto_acceso"],
         "Armario calle: 5 k€ one-shot vs 8 k€/año alquiler"],
        ["CAPEX extra A900", f"{dc['capex_extra']:,.0f} €",
         "CPD, NOC, salida internet, redundancia"],
    ]
    _add_table(doc, ["Palanca", "Valor", "Justificación"], rows,
               [5.0, 4.5, 7.5])

    _add_heading(doc, "3.1. Justificación de la estrategia rampa", level=2)
    _add_paragraph(
        doc,
        "La estrategia de cuota mayorista entre operadores admite tres "
        "modalidades: equilibrio (entramos como un operador más, α = 1/(n+1)), "
        "agresivo (desplazamos a uno, α = 1/n) o rampa (cuota objetivo "
        "alcanzada progresivamente). Hemos optado por rampa porque es la "
        "única que refleja la realidad comercial: ningún operador entra "
        "en un mercado con cuota plena el año 1, sino que la captura es "
        "progresiva conforme se firman contratos con los operadores "
        "minoristas y se demuestra calidad de servicio. Una α=0,40 "
        "objetivo a 8 años es defendible frente a 4 operadores presentes "
        "(media del territorio) y deja margen para una negociación al alza.",
    )

    _add_heading(doc, "3.2. Justificación del overbooking", level=2)
    _add_paragraph(
        doc,
        "El overbooking aplicado se corresponde con los estándares de la "
        "industria de telecomunicaciones: 1:20 residencial, 1:5 empresa "
        "y 1:1 ABAST. La diferenciación es clave porque no degradamos el "
        "BW de las sedes públicas (que requieren tráfico garantizado por "
        "criticidad operativa), pero sí aprovechamos la naturaleza "
        "estadística del consumo residencial para optimizar el "
        "dimensionamiento de la red. El escenario optimista propone "
        "1:25 residencial bajo la hipótesis de un mercado consolidado "
        "donde el patrón estadístico es más predecible.",
    )


def _block_4_geomarketing(
    doc: Document, df_viab: pd.DataFrame, df_orden: pd.DataFrame,
) -> None:
    _add_heading(doc, "4. Análisis geomarketing", level=1)
    _add_paragraph(
        doc,
        "El análisis geomarketing por municipio cumple tres funciones en "
        "la oferta. Primero, sustenta la estrategia y orden de "
        "despliegue: qué municipios primero, cuáles después, cuáles dejar "
        "para el final. Segundo, estima ingresos potenciales y payback "
        "municipal como input al modelo económico de la TAREA 7. Tercero, "
        "documenta la viabilidad por municipio aunque la cobertura sea "
        "obligatoria al 100% (un municipio no viable se conecta igual; "
        "el flag sirve para trazabilidad y para defender el equilibrio "
        "económico ante el mecanismo LCSP).",
    )

    _add_heading(doc, "4.1. Resumen de viabilidad", level=2)
    n_total = len(df_viab)
    n_viables = int(df_viab["viable"].sum())
    n_no_viables = n_total - n_viables
    ingresos_total = df_viab["ingresos_total_anuales"].sum() / 1e6
    _add_table(
        doc,
        ["Total munis", "Viables", "No viables",
         "Ingresos potenciales (M€/año)"],
        [[n_total, n_viables, n_no_viables, f"{ingresos_total:,.1f}"]],
        [3.0, 3.0, 3.0, 6.0],
    )
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        f"Del total de 900 municipios, {n_viables} resultan viables "
        f"económicamente (cumplen los umbrales mínimos: ingresos anuales "
        f"≥ 5.000 € y payback ≤ 12 años) y {n_no_viables} no lo son. "
        f"Recordamos que estos {n_no_viables} municipios deben "
        f"conectarse igual por imperativo contractual, lo que justifica "
        f"el mecanismo de reequilibrio económico previsto en el pliego. "
        f"Los ingresos potenciales totales en régimen estacionario "
        f"alcanzan {ingresos_total:,.1f} M€/año.",
    )

    _add_heading(doc, "4.2. Top 5 municipios por score de atractivo", level=2)
    top_score = df_viab.nlargest(5, "score_atractivo")[
        ["municipio", "hab", "sedes_abast", "ingresos_total_anuales",
         "payback_municipal_anios", "score_atractivo"]
    ]
    score_rows = []
    for _, r in top_score.iterrows():
        score_rows.append([
            r["municipio"],
            f"{int(r['hab']):,}",
            int(r["sedes_abast"]),
            f"{r['ingresos_total_anuales']:,.0f}",
            "∞" if r["payback_municipal_anios"] == float("inf")
            else f"{r['payback_municipal_anios']:.1f}",
            f"{r['score_atractivo']:.1f}",
        ])
    _add_table(
        doc,
        ["Muni", "Hab", "Sedes ABAST", "Ingresos €/año",
         "Payback (años)", "Score"],
        score_rows,
        [2.0, 2.5, 2.5, 4.0, 2.5, 2.0],
    )
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "El score de atractivo combina cuatro factores: ingresos anuales "
        "potenciales (peso 50%), payback corto (30%, invertido), densidad "
        "de PYMEs estimadas (10%) y baja competencia (10%, invertido). "
        "A900 sale como el municipio más atractivo del territorio (769 "
        "sedes ABAST y 2 millones de habitantes), seguido de los nodos "
        "troncales que también acumulan tráfico de los anillos de "
        "agregación que gobiernan.",
    )

    _add_heading(doc, "4.3. Plan de despliegue plurianual", level=2)
    _add_paragraph(
        doc,
        "El orden de despliegue se determina por ingresos potenciales "
        "descendentes, con el municipio inicial fijado en A900 (datacenter "
        "y centro de operaciones). El plan se rige además por dos topes "
        "anuales que reflejan el cuello de botella de obra civil y "
        "personal: 500 km de red nueva por año y 400 sedes ABAST "
        "iluminadas por año (escenario base; pesimista los baja a "
        "350/280, optimista los mantiene). Los compromisos contractuales "
        "ofrecidos son: 40% de cobertura al año 5, 80% al año 10 y 100% "
        "al año 15.",
    )

    top_orden = df_orden.head(10)
    orden_rows = []
    for _, r in top_orden.iterrows():
        orden_rows.append([
            int(r["orden_despliegue"]),
            r["municipio"],
            f"{r['criterio_valor']:,.0f}",
            int(r["anio_objetivo_despliegue"]),
        ])
    _add_table(
        doc,
        ["Orden", "Muni", "Ingresos potenciales (€/año)", "Año despliegue"],
        orden_rows,
        [2.0, 2.5, 6.0, 3.0],
    )


def _block_5_findings(
    doc: Document, df_equipment: pd.DataFrame,
    img_anillos: Path, img_a2: Path,
) -> None:
    _add_heading(doc, "5. Hallazgos clave para la defensa", level=1)
    _add_paragraph(
        doc,
        "Esta sección recoge los tres hallazgos más relevantes derivados "
        "del análisis técnico, que serán el eje argumental de la defensa "
        "oral del 12/05 y de la negociación posterior con la "
        "administración.",
    )

    bw_a900 = float(
        df_equipment.loc[df_equipment["municipio"] == "A900", "bw_acumulado_mbps"].iloc[0]
    ) / 1000
    bw_a800 = float(
        df_equipment.loc[df_equipment["municipio"] == "A800", "bw_acumulado_mbps"].iloc[0]
    ) / 1000

    _add_heading(doc, "5.1. A900 es el SPOF crítico de la red", level=2)
    _add_callout(
        doc,
        f"El 100% del tráfico territorial ({bw_a900:,.1f} Gbps) converge "
        f"en A900, que es a la vez datacenter del operador, NOC y punto "
        f"de salida a internet pública para los clientes mayoristas. "
        f"Justifica los 5 M€ extra de CAPEX específico (energía, "
        f"refrigeración, seguridad, IT) y la decisión de redundancia 1+1 "
        f"en la capa troncal y en el propio nodo A900.",
    )
    _add_paragraph(
        doc,
        "A900 está físicamente sobre los tres anillos troncales (T1, T2, "
        "T3) y los dos primeros anillos de agregación (A1, A2). Toda la "
        "topología converge ahí. Por eso se contempla un PoP secundario "
        "como opción configurable en el escenario optimista (ver "
        "decisiones.yaml, sección datacenter_a900.pop_secundario_*), con "
        "replicación síncrona en otro nodo troncal para garantizar "
        "continuidad ante un fallo total de A900. El escenario base no "
        "lo activa por contención de CAPEX, pero es una palanca clave "
        "para la negociación si la administración exige un SLA más "
        "estricto que el 99,9% comprometido.",
    )

    _add_heading(doc, "5.2. A2 es la espina dorsal del territorio", level=2)
    _add_callout(
        doc,
        f"El 88% de los municipios de acceso (793 de 900) resuelven a "
        f"través del anillo A2. Es deliberado del diseño que recibimos "
        f"de la administración. A800 (gateway de A2) acumula "
        f"{bw_a800:,.1f} Gbps y necesita MPLS + 17 cajas 40p en paralelo "
        f"+ óptico 40λ. Es el segundo punto de fallo más crítico tras "
        f"A900.",
    )
    _add_paragraph(
        doc,
        "El anillo A2 cuelga directamente de A900 y agrega la inmensa "
        "mayoría del tráfico de acceso del territorio. El gateway, A800, "
        "es el segundo nodo más crítico de toda la red y requiere un "
        "sobredimensionamiento muy notable respecto a los demás gateways "
        "de agregación (que típicamente sólo necesitan un MPLS + óptico "
        "40λ). El plan de despliegue prioriza iluminar A2 en los "
        "primeros años para asegurar la viabilidad del modelo de "
        "ingresos.",
    )

    _add_heading(
        doc, "5.3. Sensibilidad α × penetración para el reequilibrio",
        level=2,
    )
    _add_callout(
        doc,
        "El rango de ingresos entre escenarios (5,3 M€/año pesimista → "
        "7,5 M€/año optimista) está dominado por el producto "
        "cuota×penetración. La separación explícita de los dos factores "
        "que pidió el profesor el 07/05 permite atribuir el riesgo "
        "correctamente: la penetración es factor exógeno (mercado), la "
        "cuota es decisión propia (estrategia comercial). Esto es "
        "decisivo para la negociación del reequilibrio LCSP.",
    )
    _add_paragraph(
        doc,
        "El mecanismo de reequilibrio económico previsto en la LCSP se "
        "activa cuando la realidad se desvía sustancialmente del plan "
        "financiero presentado en la oferta. Si la desviación va a favor "
        "del concesionario (clawback), la administración reclama parte "
        "del extra; si va en contra, el concesionario reclama "
        "compensación. La separación entre cuota y penetración nos "
        "permite, en una hipotética negociación, atribuir las "
        "desviaciones a su causa real: si la penetración del territorio "
        "queda por debajo de lo previsto (factor de mercado, no "
        "controlable), tenemos base para reclamar compensación; si la "
        "cuota efectiva es menor (decisión comercial defectuosa), "
        "asumimos el coste. Esta asignación causal es uno de los "
        "argumentos más fuertes de la oferta.",
    )

    _add_heading(doc, "5.4. Diagrama de la topología", level=2)
    if img_anillos.exists():
        doc.add_picture(str(img_anillos), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figura 5.1 — Anillos troncal (rojo) + agregación (gris)")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    if img_a2.exists():
        doc.add_paragraph()
        doc.add_picture(str(img_a2), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figura 5.2 — Detalle del anillo A2 (espina dorsal)")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY


def _block_conclusions(doc: Document) -> None:
    _add_heading(doc, "6. Conclusiones y próximos pasos", level=1)
    _add_paragraph(
        doc,
        "La oferta técnica que se presenta cumple el 100% de los "
        "requisitos del pliego: cobertura completa de los 900 municipios "
        "y las 5.000 sedes ABAST, dimensionamiento técnico coherente con "
        "el tráfico esperado en régimen estacionario y plan de "
        "despliegue plurianual con compromisos de cobertura escalonados. "
        "El CAPEX de equipos (18,98 M€ en escenario base) es contenido y "
        "estable entre escenarios, lo que reduce el riesgo de desviación "
        "respecto al modelo económico que se presentará en la TAREA 7.",
    )
    _add_paragraph(
        doc,
        "Los próximos pasos planificados son: TAREA 6 (19/05/2026) con "
        "el diseño detallado de la solución técnica, incluyendo redundancia "
        "y SLA; entrega oficial de la oferta técnica el 26/05/2026; y "
        "TAREA 7 (02/06/2026) con el modelo económico (CAPEX, OPEX, "
        "ingresos, P&L y KPIs financieros NPV/IRR/payback) que dará "
        "lugar a la entrega de la oferta económica el 09/06/2026.",
    )
    _add_paragraph(
        doc,
        "Tras la ronda de negociación en la que la administración pedirá "
        "modificaciones concretas, la oferta final se defenderá oralmente "
        "el 30/06/2026.",
    )

    _add_heading(doc, "Anexo — Reproducibilidad", level=2)
    _add_paragraph(
        doc,
        "Todos los resultados de este documento son reproducibles "
        "ejecutando: 'python main.py --escenario todos' (recalcula los "
        "outputs CSV de los tres escenarios) y 'python -m "
        "src.oferta_tecnica_docx --escenario base' (regenera este DOCX). "
        "El archivo data/decisiones.yaml contiene todas las palancas "
        "variables; cualquier cambio recalcula automáticamente la "
        "pipeline completa en menos de 30 segundos.",
    )


# ---------------------------------------------------------------------------
# Compositor principal
# ---------------------------------------------------------------------------

def generate_oferta_tecnica_docx(
    escenario: str = "base",
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    data_dir: Path = DEFAULT_DATA_DIR,
) -> Path:
    """Genera el DOCX de oferta técnica para el escenario dado."""
    scenarios_data = {esc: _load_outputs(output_dir, esc) for esc in ESCENARIOS}
    decisiones = _load_decisiones(data_dir / "decisiones.yaml", escenario)

    out_oferta = output_dir / "oferta_tecnica"
    chart_bw = out_oferta / "bw_por_escenario.png"
    chart_capex = out_oferta / "capex_por_escenario.png"
    img_dir = output_dir / f"escenario_{escenario}" / "img"
    img_anillos = img_dir / "topologia_anillos.png"
    img_a2 = img_dir / "topologia_detalle_A2.png"

    doc = Document()

    # Estilo Normal por defecto
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Márgenes
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Bloques
    _block_cover(doc, escenario)
    _block_intro(doc)
    doc.add_page_break()
    _block_1_traffic(doc, scenarios_data, chart_bw)
    doc.add_page_break()
    _block_2_equipment(doc, scenarios_data[escenario]["equipment"], chart_capex)
    doc.add_page_break()
    _block_3_decisions(doc, decisiones)
    doc.add_page_break()
    _block_4_geomarketing(
        doc,
        scenarios_data[escenario]["viability"],
        scenarios_data[escenario]["deployment"],
    )
    doc.add_page_break()
    _block_5_findings(
        doc, scenarios_data[escenario]["equipment"], img_anillos, img_a2,
    )
    doc.add_page_break()
    _block_conclusions(doc)

    out_oferta.mkdir(parents=True, exist_ok=True)
    docx_path = out_oferta / f"oferta_tecnica_TAREA5_{escenario}.docx"
    doc.save(str(docx_path))
    logger.info("Oferta técnica DOCX generada: %s", docx_path)
    return docx_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Genera el DOCX de oferta técnica TAREA 5.",
    )
    p.add_argument(
        "--escenario", default="base", choices=ESCENARIOS,
        help="Escenario defendido (default: base).",
    )
    p.add_argument(
        "--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR,
        help="Ruta base de outputs/.",
    )
    p.add_argument(
        "--data-dir", type=Path, default=DEFAULT_DATA_DIR,
        help="Ruta a data/.",
    )
    return p.parse_args(argv)


if __name__ == "__main__":
    args = _parse_args()
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )
    docx_path = generate_oferta_tecnica_docx(
        args.escenario, args.output_dir, args.data_dir,
    )
    print(f"OK -> {docx_path}")
